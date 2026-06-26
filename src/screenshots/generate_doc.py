"""Generate Chuong3_KetQuaNghienCuu.docx based on screenshots."""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
OUT = ROOT / "Chuong3_KetQuaNghienCuu_v2.docx"

# (so_muc, tieu_de_muc, ten_anh, ten_hinh, chuc_nang, mo_ta, ket_qua)
SECTIONS = [
    ("3.1", "Giao diện trang chủ",
     "public/01_home.png", "giao diện trang chủ",
     "Trang chủ giới thiệu",
     "Trang chủ là điểm đến đầu tiên khi người dùng truy cập website. Giao diện trình bày tổng quan về dịch vụ quản lý chi tiêu, các tính năng nổi bật, lời kêu gọi hành động đăng ký/đăng nhập và phần giới thiệu đội ngũ phát triển.",
     "Người dùng nắm được giá trị cốt lõi của hệ thống và có thể nhanh chóng chuyển sang trang đăng ký hoặc đăng nhập để bắt đầu sử dụng."),

    ("3.2", "Giao diện đăng ký",
     "public/04_register.png", "giao diện đăng ký tài khoản",
     "Đăng ký tài khoản",
     "Website cung cấp giao diện để người dùng nhập email và mật khẩu. Hệ thống có cơ chế xác minh qua mã OTP gửi tới email trước khi hoàn tất đăng ký.",
     "Hệ thống kiểm tra dữ liệu hợp lệ, lưu thông tin vào cơ sở dữ liệu và tạo tài khoản mới sau khi xác minh OTP thành công. Người dùng có thể đăng nhập để sử dụng đầy đủ các tính năng của website."),

    ("3.3", "Giao diện đăng nhập",
     "public/03_login.png", "giao diện đăng nhập",
     "Đăng nhập tài khoản",
     "Form đăng nhập yêu cầu người dùng nhập email và mật khẩu đã đăng ký. Có tùy chọn ghi nhớ đăng nhập và liên kết tới chức năng quên mật khẩu.",
     "Sau khi xác thực thành công, hệ thống chuyển hướng người dùng vào trang Dashboard để bắt đầu quản lý chi tiêu cá nhân. Trường hợp sai thông tin, hệ thống hiển thị thông báo lỗi tương ứng."),

    ("3.4", "Giao diện quên mật khẩu",
     "public/05_forgot_password.png", "giao diện quên mật khẩu",
     "Khôi phục mật khẩu",
     "Người dùng nhập email đã đăng ký để nhận mã OTP đặt lại mật khẩu. Sau khi xác minh OTP, người dùng có thể nhập mật khẩu mới.",
     "Hệ thống gửi OTP về email, xác thực và cho phép người dùng đặt lại mật khẩu an toàn mà không cần truy cập tài khoản cũ."),

    ("3.5", "Giao diện chính sách bảo mật",
     "public/02_privacy.png", "giao diện chính sách bảo mật",
     "Hiển thị chính sách bảo mật",
     "Trang trình bày các điều khoản, chính sách bảo mật và cách hệ thống xử lý dữ liệu cá nhân của người dùng.",
     "Người dùng nắm rõ quyền lợi, nghĩa vụ khi sử dụng dịch vụ và yên tâm về việc bảo vệ thông tin cá nhân."),

    ("3.6", "Giao diện bảng điều khiển (Dashboard)",
     "user1/01_dashboard.png", "giao diện bảng điều khiển",
     "Tổng quan tình hình chi tiêu",
     "Dashboard hiển thị các chỉ số nhanh: tổng chi tiêu trong tháng, chi tiêu hôm nay, hạn mức còn lại, biểu đồ chi tiêu theo danh mục và danh sách giao dịch gần đây. Có khu vực gợi ý từ AI hỗ trợ người dùng quản lý tài chính.",
     "Người dùng có cái nhìn tổng quan về tình hình tài chính cá nhân ngay khi đăng nhập, nhanh chóng phát hiện bất thường và đưa ra điều chỉnh chi tiêu kịp thời."),

    ("3.7", "Giao diện danh sách chi tiêu",
     "user1/02_chitieu_index.png", "giao diện danh sách chi tiêu",
     "Quản lý chi tiêu",
     "Trang hiển thị danh sách các khoản chi tiêu của người dùng theo dạng bảng, hỗ trợ tìm kiếm theo tên, lọc theo danh mục và khoảng ngày. Mỗi dòng cho phép xem chi tiết, chỉnh sửa hoặc xóa.",
     "Người dùng dễ dàng tra cứu, theo dõi và chỉnh sửa các khoản chi tiêu đã ghi nhận, đảm bảo dữ liệu luôn chính xác."),

    ("3.8", "Giao diện thêm chi tiêu",
     "user1/03_chitieu_create.png", "giao diện thêm chi tiêu mới",
     "Thêm khoản chi tiêu",
     "Form nhập liệu cho phép người dùng nhập tên khoản chi, số tiền, ngày chi, ghi chú và chọn danh mục. Hệ thống có tính năng gợi ý danh mục tự động dựa trên tên khoản chi.",
     "Sau khi lưu, khoản chi tiêu được ghi nhận vào cơ sở dữ liệu, cập nhật ngay lập tức trên Dashboard, danh sách chi tiêu và các báo cáo thống kê."),

    ("3.9", "Giao diện danh mục chi tiêu",
     "user1/04_danhmuc.png", "giao diện quản lý danh mục",
     "Quản lý danh mục chi tiêu",
     "Trang liệt kê các danh mục chi tiêu (ăn uống, di chuyển, mua sắm…) kèm màu sắc nhận diện. Người dùng có thể thêm mới, chỉnh sửa hoặc xóa danh mục.",
     "Người dùng tổ chức được hệ thống danh mục phù hợp với thói quen chi tiêu, giúp việc phân loại và thống kê các khoản chi rõ ràng và trực quan hơn."),

    ("3.10", "Giao diện thống kê và báo cáo",
     "user1/05_thongke.png", "giao diện thống kê chi tiêu",
     "Thống kê và báo cáo",
     "Trang hiển thị các biểu đồ phân tích chi tiêu theo tháng, theo danh mục và so sánh giữa các kỳ. Người dùng có thể chọn năm để xem báo cáo chi tiết.",
     "Người dùng nhận diện được xu hướng chi tiêu, các danh mục tiêu tốn nhiều tiền nhất và tự đánh giá hiệu quả quản lý tài chính cá nhân."),

    ("3.11", "Giao diện chi tiêu định kỳ",
     "user1/06_lichchitieu_index.png", "giao diện chi tiêu định kỳ",
     "Danh sách chi tiêu theo lịch",
     "Trang quản lý các khoản chi tiêu định kỳ như tiền nhà, hóa đơn, thuê bao… Hệ thống chia thành hai khu vực: các khoản đã hoàn thành và các khoản sắp đến hạn.",
     "Người dùng theo dõi được các khoản chi cố định, không bỏ sót những khoản thanh toán quan trọng. Hệ thống tự động ghi nhận chi tiêu khi đến hạn."),

    ("3.12", "Giao diện thêm chi tiêu định kỳ",
     "user1/07_lichchitieu_create.png", "giao diện thêm chi tiêu định kỳ",
     "Thêm khoản chi tiêu định kỳ",
     "Form cho phép người dùng nhập tên khoản chi, số tiền, ngày thực hiện và danh mục. Khi đến ngày hẹn, hệ thống sẽ tự động tạo bản ghi chi tiêu tương ứng.",
     "Người dùng thiết lập được các khoản chi cố định một lần duy nhất, tiết kiệm thao tác nhập liệu hằng tháng và đảm bảo không quên các khoản chi quan trọng."),

    ("3.13", "Giao diện đặt hạn mức chi tiêu",
     "user1/08_gioihan.png", "giao diện đặt hạn mức chi tiêu",
     "Quản lý hạn mức chi tiêu",
     "Trang cho phép người dùng đặt hạn mức tối đa cho từng tháng, hiển thị tỉ lệ chi tiêu so với hạn mức đã đặt và cảnh báo khi vượt ngưỡng.",
     "Người dùng kiểm soát được ngân sách hàng tháng, nhận cảnh báo kịp thời khi sắp vượt mức cho phép, hỗ trợ duy trì kỷ luật tài chính."),

    ("3.14", "Giao diện thông báo",
     "user1/09_thongbao.png", "giao diện thông báo hệ thống",
     "Quản lý thông báo",
     "Trang liệt kê các thông báo từ hệ thống: nhắc thanh toán hóa đơn, cảnh báo vượt hạn mức, ghi nhận chi tiêu định kỳ… với trạng thái đã đọc/chưa đọc.",
     "Người dùng không bỏ sót các sự kiện quan trọng liên quan đến tài chính cá nhân, có thể đánh dấu đã đọc để theo dõi tốt hơn."),

    ("3.15", "Giao diện hồ sơ cá nhân",
     "user1/10_profile.png", "giao diện hồ sơ cá nhân",
     "Quản lý hồ sơ người dùng",
     "Trang cho phép người dùng cập nhật thông tin cá nhân: họ tên, email, ảnh đại diện. Bên cạnh đó cung cấp chức năng đổi mật khẩu (xác minh qua OTP) và xóa tài khoản.",
     "Người dùng chủ động cập nhật thông tin tài khoản, đảm bảo an toàn bảo mật thông qua xác thực hai bước khi thực hiện các thao tác nhạy cảm."),

    ("3.16", "Giao diện cài đặt nhắc nhở",
     "user1/11_settings.png", "giao diện cài đặt nhắc nhở",
     "Cài đặt nhắc nhở qua email",
     "Trang cho phép người dùng bật/tắt nhận email nhắc nhở, chọn tần suất (hằng ngày, hằng tuần, hằng tháng) và giờ nhận thông báo.",
     "Người dùng được nhắc nhở ghi chép chi tiêu đều đặn theo lịch cá nhân hóa, hình thành thói quen quản lý tài chính một cách nhất quán."),

    ("3.17", "Giao diện quản lý người dùng (dành cho Admin)",
     "admin/12_admin_users.png", "giao diện quản lý người dùng",
     "Quản lý danh sách người dùng",
     "Trang chỉ hiển thị với tài khoản có vai trò Admin. Giao diện trình bày các chỉ số tổng quan (tổng số người dùng, số quản trị viên, số người dùng thường), thanh tìm kiếm theo email/họ tên và bảng danh sách kèm avatar, vai trò, số lượng và tổng chi tiêu của mỗi người dùng.",
     "Quản trị viên có cái nhìn tổng thể về người dùng đang sử dụng hệ thống, dễ dàng tìm kiếm và thực hiện các thao tác sửa, xóa, đổi vai trò trực tiếp trên cùng một trang."),

    ("3.18", "Giao diện thêm người dùng (dành cho Admin)",
     "admin/13_admin_create.png", "giao diện thêm người dùng",
     "Thêm tài khoản người dùng",
     "Form cho phép quản trị viên tạo tài khoản mới với các thông tin: email, họ tên, mật khẩu, xác nhận mật khẩu và vai trò (Admin hoặc User). Hệ thống có kiểm tra trùng email và độ khớp của hai trường mật khẩu.",
     "Sau khi lưu thành công, tài khoản mới được tạo và xuất hiện ngay trong danh sách người dùng. Quản trị viên có thể nhanh chóng cấp tài khoản cho thành viên mới mà không cần thao tác đăng ký công khai."),

    ("3.19", "Giao diện chỉnh sửa người dùng (dành cho Admin)",
     "admin/14_admin_edit.png", "giao diện chỉnh sửa người dùng",
     "Cập nhật tài khoản người dùng",
     "Form hiển thị thông tin hiện tại của tài khoản, cho phép cập nhật email, họ tên, vai trò và đặt lại mật khẩu (để trống nếu không đổi). Trường vai trò bị khóa khi quản trị viên đang chỉnh sửa chính tài khoản của mình để tránh mất quyền truy cập.",
     "Quản trị viên cập nhật được thông tin tài khoản, nâng/hạ quyền hoặc đặt lại mật khẩu cho người dùng khi cần thiết, hỗ trợ vận hành hệ thống một cách linh hoạt."),
]


def add_centered(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_section(doc, idx_hinh, so_muc, tieu_de, anh_path, ten_hinh, chuc_nang, mo_ta, ket_qua):
    h = doc.add_heading(level=2)
    h.add_run(f"{so_muc}. {tieu_de}")

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(ROOT / anh_path), width=Cm(15))

    add_centered(doc, f"Hình {idx_hinh} {ten_hinh}", size=11)

    p1 = doc.add_paragraph()
    p1.add_run("Chức năng: ").bold = True
    p1.add_run(chuc_nang)

    p2 = doc.add_paragraph()
    p2.add_run("Mô tả: ").bold = True
    p2.add_run(mo_ta)

    p3 = doc.add_paragraph()
    p3.add_run("Kết quả: ").bold = True
    p3.add_run(ket_qua)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("CHƯƠNG 3\tKẾT QUẢ NGHIÊN CỨU")

    start_hinh = 24
    for i, sec in enumerate(SECTIONS):
        add_section(doc, start_hinh + i, *sec)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
